"""
Structured SRS document parser.
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Tuple

import docx
from pypdf import PdfReader

from core.schemas import EvidenceSpan, ParsedSection, ParsedSRSDocument, RequirementRecord


class SRSDocumentParser:
    """Parse SRS documents into clean text plus structured sections/requirements."""

    def __init__(self):
        self.toc_pattern = re.compile(r"\.{4,}\s*\d+")
        self.header_footer_pattern = re.compile(r"^\s*\d+\s*$|^\s*page\s*\d+\s*$", re.IGNORECASE)
        self.functional_requirement_pattern = re.compile(r"^functional requirement\s+\d+", re.IGNORECASE)
        self.user_class_pattern = re.compile(r"^user class\s+\d+\s*[-:]\s*(.+)$", re.IGNORECASE)
        self.known_headings = {
            "introduction",
            "purpose",
            "scope",
            "problem definition",
            "audience",
            "tools to be used",
            "references",
            "overview",
            "glossary",
            "definitions",
            "acronyms and abbreviations",
            "system model",
            "functional requirement",
            "functional requirements",
            "non functional requirements",
            "non-functional requirements",
            "security requirements",
            "performance requirements",
            "safety requirements",
            "system evolution",
            "requirement specification",
            "product perspective",
            "hardware interface",
            "memory constraint",
            "product function",
            "database requirement",
            "user characteristics",
            "assumptions and dependencies",
        }

    def parse_file(self, file_path: str) -> str:
        return self.parse_structured_file(file_path).clean_text

    def parse_structured_file(self, file_path: str) -> ParsedSRSDocument:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(file_path)

        ext = path.suffix.lower()
        if ext == ".pdf":
            raw_lines = self._read_pdf_lines(file_path)
        elif ext == ".docx":
            raw_lines = self._read_docx_lines(file_path)
        elif ext == ".txt":
            raw_lines = self._read_txt_lines(file_path)
        else:
            raise Exception(f"Unsupported file type: {ext}")

        clean_lines = self._clean_lines(raw_lines)
        clean_lines = self._truncate_at_references(clean_lines)
        sections, evidence_spans = self._build_sections(path.name, clean_lines)
        requirements = self._extract_requirement_records(sections, evidence_spans)
        clean_text = "\n".join(line for _, line in clean_lines)

        return ParsedSRSDocument(
            file_path=str(path),
            document_name=path.name,
            clean_text=clean_text,
            sections=sections,
            requirements=requirements,
            evidence_spans=evidence_spans,
        )

    def _truncate_at_references(self, lines: List[Tuple[int, str]]) -> List[Tuple[int, str]]:
        stop_words = {"references", "bibliography", "kaynakça"}
        for idx, (_, line) in enumerate(lines):
            if line.strip().lower() in stop_words:
                return lines[:idx]
        return lines

    def _clean_lines(self, raw_lines: List[str]) -> List[Tuple[int, str]]:
        cleaned: List[Tuple[int, str]] = []
        for raw_index, line in enumerate(raw_lines, start=1):
            stripped = line.strip()
            if not stripped:
                continue
            if self.header_footer_pattern.match(stripped):
                continue
            if self.toc_pattern.search(stripped):
                continue
            cleaned.append((raw_index, stripped))
        return cleaned

    def _build_sections(
        self,
        document_name: str,
        clean_lines: List[Tuple[int, str]],
    ) -> tuple[List[ParsedSection], List[EvidenceSpan]]:
        sections: List[ParsedSection] = []
        evidence_spans: List[EvidenceSpan] = []
        current_heading = "Document"
        current_category = "general"
        current_lines: List[str] = []
        current_evidence_ids: List[str] = []
        section_counter = 1
        evidence_counter = 1

        def flush_section() -> None:
            nonlocal section_counter, current_lines, current_evidence_ids
            if not current_lines:
                return
            sections.append(
                ParsedSection(
                    section_id=f"SEC-{section_counter:03d}",
                    heading=current_heading,
                    category=current_category,
                    content="\n".join(current_lines).strip(),
                    evidence_ids=list(current_evidence_ids),
                )
            )
            section_counter += 1
            current_lines = []
            current_evidence_ids = []

        for line_number, text in clean_lines:
            if self._is_heading(text):
                flush_section()
                current_heading = text
                current_category = self._classify_heading(text)
                continue

            evidence_id = f"E{evidence_counter:04d}"
            evidence_counter += 1
            current_lines.append(text)
            current_evidence_ids.append(evidence_id)
            evidence_spans.append(
                EvidenceSpan(
                    evidence_id=evidence_id,
                    document=document_name,
                    section=current_heading,
                    excerpt=text,
                    line=line_number,
                )
            )

        flush_section()
        return sections, evidence_spans

    def _extract_requirement_records(
        self,
        sections: List[ParsedSection],
        evidence_spans: List[EvidenceSpan],
    ) -> List[RequirementRecord]:
        evidence_by_id = {e.evidence_id: e for e in evidence_spans}
        records: List[RequirementRecord] = []
        actor_context: str | None = None
        record_counter = 1

        for section in sections:
            lines = [line.strip() for line in section.content.splitlines() if line.strip()]
            if not lines:
                continue

            if section.category == "functional":
                idx = 0
                while idx < len(lines):
                    user_match = self.user_class_pattern.match(lines[idx])
                    if user_match:
                        actor_context = self._normalize_actor(user_match.group(1))
                        idx += 1
                        continue

                    if self.functional_requirement_pattern.match(lines[idx]):
                        title = ""
                        description_parts: List[str] = []
                        block_lines = [lines[idx]]
                        idx += 1
                        while idx < len(lines) and not self.functional_requirement_pattern.match(lines[idx]) and not self.user_class_pattern.match(lines[idx]):
                            block_lines.append(lines[idx])
                            if lines[idx].lower().startswith("title:"):
                                title = lines[idx].split(":", 1)[1].strip()
                            elif lines[idx].lower().startswith("desc:"):
                                description_parts.append(lines[idx].split(":", 1)[1].strip())
                            else:
                                description_parts.append(lines[idx])
                            idx += 1

                        evidence_ids = self._match_evidence_ids(block_lines, section.evidence_ids, evidence_by_id)
                        records.append(
                            RequirementRecord(
                                requirement_id=f"REQ-{record_counter:03d}",
                                category="functional",
                                title=title or lines[idx - 1],
                                description=" ".join(part for part in description_parts if part).strip() or title or lines[idx - 1],
                                actor=actor_context,
                                section=section.heading,
                                evidence_ids=evidence_ids,
                            )
                        )
                        self._annotate_requirement_ids(records[-1], evidence_ids, evidence_by_id)
                        record_counter += 1
                        continue
                    idx += 1
                continue

            if section.category in {"security", "performance", "safety", "non_functional"}:
                for line in lines:
                    evidence_ids = self._match_evidence_ids([line], section.evidence_ids, evidence_by_id)
                    records.append(
                        RequirementRecord(
                            requirement_id=f"REQ-{record_counter:03d}",
                            category=section.category,
                            title=line[:80],
                            description=line,
                            actor=None,
                            section=section.heading,
                            evidence_ids=evidence_ids,
                        )
                    )
                    self._annotate_requirement_ids(records[-1], evidence_ids, evidence_by_id)
                    record_counter += 1
                continue

            if section.category == "database":
                idx = 0
                while idx < len(lines):
                    title = lines[idx]
                    description = lines[idx + 1] if idx + 1 < len(lines) else lines[idx]
                    block = [title, description]
                    evidence_ids = self._match_evidence_ids(block, section.evidence_ids, evidence_by_id)
                    records.append(
                        RequirementRecord(
                            requirement_id=f"REQ-{record_counter:03d}",
                            category="database",
                            title=title,
                            description=description,
                            actor=None,
                            section=section.heading,
                            evidence_ids=evidence_ids,
                        )
                    )
                    self._annotate_requirement_ids(records[-1], evidence_ids, evidence_by_id)
                    record_counter += 1
                    idx += 2
                continue

            if section.category == "glossary":
                idx = 0
                while idx < len(lines):
                    title = lines[idx]
                    description = lines[idx + 1] if idx + 1 < len(lines) else ""
                    if description:
                        evidence_ids = self._match_evidence_ids([title, description], section.evidence_ids, evidence_by_id)
                        records.append(
                            RequirementRecord(
                                requirement_id=f"REQ-{record_counter:03d}",
                                category="glossary",
                                title=title,
                                description=description,
                                actor=None,
                                section=section.heading,
                                evidence_ids=evidence_ids,
                            )
                        )
                        self._annotate_requirement_ids(records[-1], evidence_ids, evidence_by_id)
                        record_counter += 1
                        idx += 2
                    else:
                        idx += 1
                continue

            if section.category in {"product_function", "assumptions"}:
                for line in lines:
                    evidence_ids = self._match_evidence_ids([line], section.evidence_ids, evidence_by_id)
                    records.append(
                        RequirementRecord(
                            requirement_id=f"REQ-{record_counter:03d}",
                            category=section.category,
                            title=line[:80],
                            description=line,
                            actor=self._infer_actor_from_text(line),
                            section=section.heading,
                            evidence_ids=evidence_ids,
                        )
                    )
                    self._annotate_requirement_ids(records[-1], evidence_ids, evidence_by_id)
                    record_counter += 1

        return records

    def _annotate_requirement_ids(
        self,
        record: RequirementRecord,
        evidence_ids: List[str],
        evidence_by_id: dict[str, EvidenceSpan],
    ) -> None:
        for evidence_id in evidence_ids:
            if evidence_id in evidence_by_id:
                evidence_by_id[evidence_id].requirement_id = record.requirement_id

    def _match_evidence_ids(
        self,
        target_lines: List[str],
        candidate_ids: List[str],
        evidence_by_id: dict[str, EvidenceSpan],
    ) -> List[str]:
        matched: List[str] = []
        normalized_targets = [line.lower() for line in target_lines if line]
        for evidence_id in candidate_ids:
            evidence = evidence_by_id[evidence_id]
            excerpt = evidence.excerpt.lower()
            if any(target in excerpt or excerpt in target for target in normalized_targets):
                matched.append(evidence_id)
        return matched or list(candidate_ids[:1])

    def _infer_actor_from_text(self, text: str) -> str | None:
        lowered = text.lower()
        if "non-registered" in lowered or "normal users" in lowered:
            return "NonRegisteredUser"
        if "registered users" in lowered or "authorized users" in lowered:
            return "RegisteredUser"
        if "user" in lowered:
            return "User"
        return None

    def _normalize_actor(self, raw_actor: str) -> str:
        cleaned = re.sub(r"[^A-Za-z0-9]+", " ", raw_actor).strip()
        return "".join(part.capitalize() for part in cleaned.split()) or "User"

    def _classify_heading(self, heading: str) -> str:
        lowered = heading.lower()
        if "security requirements" in lowered:
            return "security"
        if "performance requirements" in lowered:
            return "performance"
        if "safety requirements" in lowered:
            return "safety"
        if "non functional requirements" in lowered or "non-functional requirements" in lowered:
            return "non_functional"
        if "functional requirement" in lowered or "user class" in lowered:
            return "functional"
        if "database requirement" in lowered:
            return "database"
        if lowered in {"definitions", "glossary", "acronyms and abbreviations"}:
            return "glossary"
        if "product function" in lowered:
            return "product_function"
        if "assumptions and dependencies" in lowered:
            return "assumptions"
        return "general"

    def _is_heading(self, text: str) -> bool:
        stripped = text.strip().strip(":")
        if not stripped or len(stripped) > 80 or stripped.endswith("."):
            return False
        lowered = stripped.lower()
        if lowered in self.known_headings:
            return True
        alpha_chars = [char for char in stripped if char.isalpha()]
        if not alpha_chars:
            return False
        uppercase_ratio = sum(char.isupper() for char in alpha_chars) / len(alpha_chars)
        return uppercase_ratio >= 0.8

    def _read_pdf_lines(self, file_path: str) -> List[str]:
        reader = PdfReader(file_path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text.splitlines()

    def _read_docx_lines(self, file_path: str) -> List[str]:
        document = docx.Document(file_path)
        return [paragraph.text for paragraph in document.paragraphs]

    def _read_txt_lines(self, file_path: str) -> List[str]:
        with open(file_path, "r", encoding="utf-8") as handle:
            return handle.read().splitlines()
