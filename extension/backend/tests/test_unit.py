"""
Unit Tests for Backend Components

These tests don't require running server or API key.
They test individual components in isolation.

Run with: pytest tests/test_unit.py -v
"""

import pytest
import json
import sys
import os
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent.parent))
sys.path.insert(0, str(Path(__file__).resolve().parents[3] / "experiments"))


def _build_sample_srs_docx(tmp_path: Path) -> Path:
    from docx import Document

    doc = Document()
    for line in [
        "INTRODUCTION",
        "Kinmail is an e-commerce platform for buyers and sellers.",
        "FUNCTIONAL REQUIREMENT",
        "User Class 1- Non Registered Users",
        "Functional Requirement 1",
        "TITLE: Register to the website",
        "DESC: Non-registered users can register and verify email before entering the platform.",
        "Functional Requirement 2",
        "TITLE: Search for products",
        "DESC: Users can search through the category of products.",
        "User Class 2- Registered Users",
        "Functional Requirement 1",
        "TITLE: Add product",
        "DESC: Only registered users can add products with price details and product condition.",
        "Functional Requirement 2",
        "TITLE: View product dashboard",
        "DESC: Registered users can view their dashboard.",
        "NON FUNCTIONAL REQUIREMENTS",
        "Security Requirements",
        "Users need to verify their email after registering to enter the platform as registered user.",
        "Only the registered users can add products.",
        "Database Requirement",
        "Users",
        "This table stores user credentials and profile details.",
        "Products",
        "This table stores product information, price and category.",
        "Offers",
        "This table stores buyer and seller offers.",
    ]:
        doc.add_paragraph(line)

    path = tmp_path / "sample_srs.docx"
    doc.save(path)
    return path


# =============================================================================
# TOKEN TRACKER TESTS
# =============================================================================

class TestTokenTracker:
    """Test TokenTracker functionality."""
    
    def test_pricing_constants(self):
        """Test pricing configuration is loaded from config."""
        from config import PricingConfig

        assert PricingConfig.MODEL_PRICING["gemini-2.5-flash"]["input_per_1m_tokens"] == 0.30
        assert PricingConfig.MODEL_PRICING["gemini-2.5-flash"]["output_per_1m_tokens"] == 2.50
        assert PricingConfig.MODEL_PRICING["gemini-2.5-flash-lite"]["input_per_1m_tokens"] == 0.10
        assert PricingConfig.MODEL_PRICING["gemini-2.5-flash-lite"]["output_per_1m_tokens"] == 0.40
    
    def test_tracker_singleton(self):
        """Test TokenTracker singleton pattern."""
        from core.token_tracker import TokenTracker
        
        # Reset first
        TokenTracker.reset()
        
        tracker1 = TokenTracker.get_instance()
        tracker2 = TokenTracker.get_instance()
        
        assert tracker1 is tracker2
    
    def test_calculate_call_cost_flash(self):
        """Test cost calculation for a configured model."""
        from core.token_tracker import TokenTracker
        
        TokenTracker.reset()
        tracker = TokenTracker.get_instance()
        
        prompt_tokens = 1000
        completion_tokens = 500
        
        cost = tracker._calculate_call_cost(
            model="gemini-2.5-flash",
            prompt_tokens=prompt_tokens,
            completion_tokens=completion_tokens,
        )
        
        expected = (prompt_tokens * (0.30 / 1_000_000) +
                   completion_tokens * (2.50 / 1_000_000))
        
        assert abs(cost - expected) < 0.000001
    
    def test_calculate_call_cost_flash_lite(self):
        """Test cost calculation for flash-lite model."""
        from core.token_tracker import TokenTracker
        
        TokenTracker.reset()
        tracker = TokenTracker.get_instance()
        
        prompt_tokens = 1000
        completion_tokens = 500
        
        cost = tracker._calculate_call_cost(
            model="gemini-2.5-flash-lite",
            prompt_tokens=prompt_tokens,
            completion_tokens=completion_tokens,
        )
        
        expected = (prompt_tokens * (0.10 / 1_000_000) +
                   completion_tokens * (0.40 / 1_000_000))
        
        assert abs(cost - expected) < 0.000001
    
    def test_track_usage_and_delta(self):
        """Test usage tracking and snapshot deltas."""
        from core.token_tracker import TokenTracker
        
        TokenTracker.reset()
        tracker = TokenTracker.get_instance()
        snapshot = tracker.snapshot()

        tracker.track_usage(
            provider="gemini",
            model="gemini-2.5-flash-lite",
            stage="Validator",
            operation="validate_advanced_checks",
            prompt_tokens=1200,
            completion_tokens=300,
            cached_tokens=200,
            parse_success=True,
            retry_count=1,
        )

        delta = tracker.delta(snapshot)
        assert delta["llm_input_tokens"] == 1000
        assert delta["llm_output_tokens"] == 300
        assert delta["cached_tokens"] == 200
        assert delta["api_calls"] == 1
        assert delta["parseable_outputs"] == 1


# =============================================================================
# VALIDATION METRICS TRACKER TESTS
# =============================================================================

class TestValidationMetricsTracker:
    """Test ValidationMetricsTracker functionality."""
    
    def test_tracker_singleton(self):
        """Test ValidationMetricsTracker singleton pattern."""
        from core.validation_metrics import ValidationMetricsTracker
        
        tracker1 = ValidationMetricsTracker.get_instance()
        tracker2 = ValidationMetricsTracker.get_instance()
        
        assert tracker1 is tracker2
    
    def test_track_validation(self):
        """Test tracking a validation."""
        from core.validation_metrics import ValidationMetricsTracker
        
        tracker = ValidationMetricsTracker.get_instance()
        initial_count = tracker.stats.total_validations
        
        tracker.track_validation(
            filename="test.py",
            file_size_chars=100,
            code_file_tokens=50,
            validation_time_ms=50.0,
            violations=[
                {"type": "SynonymViolation", "message": "test"},
                {"type": "BannedTermViolation", "message": "test"}
            ],
            has_sources=True
        )
        
        assert tracker.stats.total_validations == initial_count + 1
    
    def test_track_validation_with_stage_metrics(self):
        from core.validation_metrics import ValidationMetricsTracker

        ValidationMetricsTracker.reset()
        tracker = ValidationMetricsTracker.get_instance()
        tracker.track_validation(
            filename="test.py",
            file_size_chars=200,
            code_file_tokens=75,
            validation_time_ms=12.5,
            violations=[],
            mode="naive",
            provider="static-json",
            model="static-model",
            stage_latencies_ms={"naive_llm": 10.0, "total": 12.5},
            llm_input_tokens=50,
            llm_output_tokens=10,
            llm_total_tokens=60,
            cost_usd=0.001,
            api_calls=1,
            parseable_outputs=1,
            unparseable_outputs=0,
        )
        report = tracker.get_report(detailed=False)
        assert report["summary"]["validation_modes"]["naive"] == 1
        assert report["performance"]["avg_stage_latencies_ms"]["naive_llm"] == 10.0
        assert report["llm_usage"]["parseable_output_rate_percent"] == 100.0


# =============================================================================
# CODE PARSER TESTS
# =============================================================================

class TestCodeParser:
    """Test CodeParser functionality."""
    
    def test_parse_valid_code(self):
        """Test parsing valid Python code."""
        from core.parser import CodeParser
        
        parser = CodeParser()
        code = '''
class Order:
    def __init__(self, order_id):
        self.order_id = order_id
'''
        
        result = parser.parse_code(code, "order.py")
        
        assert "error" not in result
        assert "classes" in result
        assert len(result["classes"]) == 1
        assert result["classes"][0]["name"] == "Order"
    
    def test_parse_syntax_error(self):
        """Test parsing code with syntax error."""
        from core.parser import CodeParser
        
        parser = CodeParser()
        code = '''
def broken(
    print("missing parenthesis"
'''
        
        result = parser.parse_code(code, "broken.py")
        
        assert "error" in result
    
    def test_parse_functions(self):
        """Test parsing functions."""
        from core.parser import CodeParser
        
        parser = CodeParser()
        code = '''
def process_order(order_id: str) -> dict:
    """Process an order."""
    return {"id": order_id}
'''
        
        result = parser.parse_code(code, "service.py")
        
        assert "functions" in result
        assert len(result["functions"]) == 1
        assert result["functions"][0]["name"] == "process_order"


# =============================================================================
# DETERMINISTIC VALIDATION TESTS
# =============================================================================

class TestDeterministicValidation:
    """Test non-LLM deterministic violation detection."""

    def test_rule_based_name_violations_detects_synonyms_and_banned_terms(self):
        from core.llm_client import LLMClient

        llm = object.__new__(LLMClient)

        ast_data = {
            "filename": "sample.py",
            "classes": [
                {"name": "ClientManager"},
                {"name": "PaymentRecord"},
                {"name": "DataHelper"},
            ],
            "functions": [],
        }
        domain_rules = {
            "bounded_contexts": [
                {
                    "ubiquitous_language": {
                        "entities": [
                            {
                                "name": "Customer",
                                "synonyms_to_avoid": ["Client", "User"],
                            },
                            {
                                "name": "Payment",
                                "synonyms_to_avoid": ["PaymentRecord"],
                            },
                        ]
                    }
                }
            ],
            "global_rules": {
                "banned_global_terms": ["Manager", "Data", "Helper"]
            },
        }

        violations = llm.rule_based_name_violations(ast_data, domain_rules)
        types = [v["type"] for v in violations]

        assert "SynonymViolation" in types
        assert "BannedTermViolation" in types
        assert any("ClientManager" in v["message"] for v in violations)
        assert any("PaymentRecord" in v["message"] for v in violations)
        assert any("DataHelper" in v["message"] for v in violations)


# =============================================================================
# PROVIDER / VALIDATION SERVICE TESTS
# =============================================================================

class TestProviderAndValidationService:
    def test_architect_completion_accepts_enum_style_stop(self):
        from core.architect import DomainArchitect

        architect = object.__new__(DomainArchitect)

        class Response:
            finish_reason = "FinishReason.STOP"

        assert architect._check_response_completion(Response(), retry=0) is True

    def test_static_provider_returns_structured_result(self):
        from core.llm_provider import StaticJSONProvider
        from core.llm_client import ValidationResponse

        provider = StaticJSONProvider(
            responses={
                "validate_advanced_checks": {
                    "is_violation": False,
                    "violations": [],
                }
            }
        )
        result = provider.generate_json(
            model="static-model",
            prompt="test",
            stage="Validator",
            operation="validate_advanced_checks",
            response_schema=ValidationResponse,
        )

        assert result.parse_success is True
        assert result.finish_reason == "STOP"
        assert result.parsed.is_violation is False

    def test_validation_service_pipeline_records_stage_metrics(self):
        from core.llm_provider import StaticJSONProvider
        from core.llm_client import LLMClient
        from core.validation_service import ValidationService
        from core.validation_metrics import ValidationMetricsTracker
        from core.token_tracker import TokenTracker

        TokenTracker.reset()
        ValidationMetricsTracker.reset()
        provider = StaticJSONProvider(
            responses={
                "validate_advanced_checks": {
                    "is_violation": False,
                    "violations": [],
                }
            }
        )
        client = LLMClient(provider=provider, model_name="static-model")
        service = ValidationService(llm_client=client)
        result = service.validate(
            filename="sample.py",
            content="class ClientManager:\n    pass\n",
            domain_rules={
                "bounded_contexts": [
                    {
                        "ubiquitous_language": {
                            "entities": [{"name": "Customer", "synonyms_to_avoid": ["Client"]}],
                            "value_objects": [],
                            "domain_events": []
                        },
                        "allowed_dependencies": []
                    }
                ],
                "global_rules": {"banned_global_terms": ["Manager"]},
            },
            rag=None,
            mode="pipeline",
        )
        assert result["metrics"]["stage_latencies_ms"]["ast_parse"] >= 0
        assert result["metrics"]["stage_latencies_ms"]["deterministic_rules"] >= 0
        assert result["metrics"]["llm_total_tokens"] >= 0
        assert result["is_violation"] is True

    def test_validation_service_naive_mode_uses_naive_llm_stage(self):
        from core.llm_provider import StaticJSONProvider
        from core.llm_client import LLMClient
        from core.validation_service import ValidationService

        provider = StaticJSONProvider(
            responses={
                "validate_naive_baseline": {
                    "is_violation": True,
                    "violations": [
                        {
                            "type": "NamingConventionViolation",
                            "message": "Class 'ClientManager' may not reflect domain-driven terminology.",
                            "suggestion": "Review names."
                        }
                    ],
                }
            }
        )
        client = LLMClient(provider=provider, model_name="static-model")
        service = ValidationService(llm_client=client)
        result = service.validate(
            filename="sample.py",
            content="class ClientManager:\n    pass\n",
            domain_rules={},
            rag=None,
            mode="naive",
        )
        assert result["mode"] == "naive"
        assert result["metrics"]["stage_latencies_ms"]["naive_llm"] >= 0
        assert result["metrics"]["api_calls"] == 1


# =============================================================================
# AST MODEL SIGNAL TESTS
# =============================================================================

class TestASTModelSignals:
    """Test AST candidate extraction and model enrichment."""

    def test_extract_candidates_with_traceability(self, tmp_path):
        from core.ast_model_signals import ASTModelSignalExtractor

        sample = tmp_path / "domain_model.py"
        sample.write_text(
            """
from dataclasses import dataclass

class Order:
    def __init__(self, order_id, total):
        self.order_id = order_id
        self.total = total

    def add_item(self, item):
        return item

@dataclass
class Money:
    amount: float
    currency: str

class OrderService:
    def place(self):
        return True

class OrderAggregate:
    def __init__(self):
        self.items = []

    def add_item(self, item):
        self.items.append(item)
""",
            encoding="utf-8",
        )

        extractor = ASTModelSignalExtractor()
        candidates = extractor.extract_candidates([str(sample)], grounding_docs=[])

        assert any(c["name"] == "Order" for c in candidates["entities"])
        assert any(c["name"] == "Money" for c in candidates["value_objects"])
        assert any(c["name"] == "OrderService" for c in candidates["services"])
        assert any(c["name"] == "OrderAggregate" for c in candidates["aggregates"])

        order_entity = next(c for c in candidates["entities"] if c["name"] == "Order")
        assert 0.0 <= order_entity["confidence"] <= 1.0
        assert len(order_entity["sources"]) >= 1
        assert "file" in order_entity["sources"][0]
        assert "line" in order_entity["sources"][0]
        assert "rule" in order_entity["sources"][0]

    def test_enrich_domain_model_adds_confidence_and_sources(self, tmp_path):
        from core.ast_model_signals import ASTModelSignalExtractor
        from core.schemas import DomainModel, ProjectMetadata

        workspace = tmp_path / "workspace"
        workspace.mkdir()
        (workspace / "billing.py").write_text(
            """
class Invoice:
    def __init__(self, invoice_id, amount):
        self.invoice_id = invoice_id
        self.amount = amount
""",
            encoding="utf-8",
        )

        model = DomainModel(
            project_name="Test",
            project_metadata=ProjectMetadata(
                version="1.0.0",
                generated_at="2026-02-16",
                description="test"
            ),
            bounded_contexts=[],
            global_rules=None,
        )

        extractor = ASTModelSignalExtractor()
        enriched = extractor.enrich_domain_model(
            model,
            str(workspace),
            srs_docs=[{"path": "srs.txt", "content": "Invoice must include amount."}],
        )

        assert len(enriched.bounded_contexts) >= 1
        ul = enriched.bounded_contexts[0].ubiquitous_language
        assert ul.entities is not None
        assert any(e.name == "Invoice" for e in ul.entities)
        invoice = next(e for e in ul.entities if e.name == "Invoice")
        assert invoice.confidence >= 0.5
        assert len(invoice.sources) >= 1


# =============================================================================
# SCHEMA TESTS
# =============================================================================

class TestSchemas:
    """Test domain model schemas."""
    
    def test_domain_model_structure(self):
        """Test DomainModel schema structure."""
        from core.schemas import DomainModel, ProjectMetadata, BoundedContext, UbiquitousLanguage
        
        # Test with complete required data
        model = DomainModel(
            project_name="Test Project",
            project_metadata=ProjectMetadata(
                version="1.0.0",
                generated_at="2026-01-21T12:00:00"
            ),
            bounded_contexts=[],
            global_rules=None
        )
        
        model_dict = model.model_dump()
        
        assert "project_name" in model_dict
        assert "project_metadata" in model_dict
        assert "bounded_contexts" in model_dict
    
    def test_bounded_context_schema(self):
        """Test BoundedContext schema."""
        from core.schemas import BoundedContext, UbiquitousLanguage, Entity
        
        context = BoundedContext(
            context_name="Order Processing",
            description="Handles order lifecycle",
            ubiquitous_language=UbiquitousLanguage(
                entities=[
                    Entity(name="Order", description="A purchase request"),
                    Entity(name="OrderItem", description="A line item in an order")
                ],
                value_objects=[],
                domain_events=[]
            )
        )
        
        assert context.context_name == "Order Processing"
        assert len(context.ubiquitous_language.entities) == 2


# =============================================================================
# DOCUMENT PARSER TESTS
# =============================================================================

class TestDocumentParser:
    """Test SRSDocumentParser functionality."""
    
    def test_parse_txt_file(self, tmp_path):
        """Test parsing a .txt file."""
        from core.document_parser import SRSDocumentParser
        
        # Create temp file
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("# Test SRS\n\nThis is a test document.")
        
        parser = SRSDocumentParser()
        content = parser.parse_file(str(txt_file))
        
        assert "Test SRS" in content
        assert "test document" in content
    
    def test_parse_nonexistent_file(self):
        """Test parsing a non-existent file raises error."""
        from core.document_parser import SRSDocumentParser
        
        parser = SRSDocumentParser()
        
        with pytest.raises(FileNotFoundError):
            parser.parse_file("/nonexistent/file.txt")

    def test_parse_structured_docx_extracts_requirements(self, tmp_path):
        from core.document_parser import SRSDocumentParser

        srs_path = _build_sample_srs_docx(tmp_path)
        parser = SRSDocumentParser()
        parsed = parser.parse_structured_file(str(srs_path))

        assert parsed.document_name == "sample_srs.docx"
        assert len(parsed.sections) >= 3
        assert any(record.actor == "NonRegisteredUsers" for record in parsed.requirements)
        assert any(record.title == "Register to the website" for record in parsed.requirements)
        assert any(record.category == "security" for record in parsed.requirements)
        assert any(record.category == "database" and record.title == "Users" for record in parsed.requirements)


# =============================================================================
# DOMAIN GENERATION PIPELINE TESTS
# =============================================================================

class TestDomainGenerationPipeline:
    def test_architect_pipeline_preserves_rich_fields(self, tmp_path):
        from core.architect import DomainArchitect
        from core.document_parser import SRSDocumentParser
        from core.llm_provider import StaticJSONProvider

        srs_path = _build_sample_srs_docx(tmp_path)
        parsed = SRSDocumentParser().parse_structured_file(str(srs_path))

        provider = StaticJSONProvider(
            responses={
                "Scout": {
                    "requirements": [],
                    "actors": [],
                    "entities": [],
                    "constraints": [],
                    "tables": [],
                    "capabilities": [],
                    "evidence_spans": [],
                },
                "Architect": {
                    "contexts": [
                        {
                            "context_name": "IdentityAndAccess",
                            "description": "Registration and access control.",
                            "ownership_rationale": "Owns registration and identity constraints.",
                            "included_capabilities": ["RegisterAccount"],
                            "excluded_capabilities": [],
                            "primary_entities": ["User"],
                            "allowed_dependencies": [],
                            "evidence_ids": ["E0001"],
                        },
                        {
                            "context_name": "SellerWorkspace",
                            "description": "Seller-side listing and dashboard features.",
                            "ownership_rationale": "Owns seller-only capabilities.",
                            "included_capabilities": ["AddProduct", "ViewProductDashboard"],
                            "excluded_capabilities": [],
                            "primary_entities": ["Product", "Offer"],
                            "allowed_dependencies": ["IdentityAndAccess", "ProductCatalog"],
                            "evidence_ids": ["E0002"],
                        },
                        {
                            "context_name": "ProductCatalog",
                            "description": "Search and browse product information.",
                            "ownership_rationale": "Owns searchable product facts and categories.",
                            "included_capabilities": ["SearchProducts", "ViewProductDetails"],
                            "excluded_capabilities": [],
                            "primary_entities": ["Product", "Category"],
                            "allowed_dependencies": [],
                            "evidence_ids": ["E0003"],
                        },
                        {
                            "context_name": "NegotiationAndCommunication",
                            "description": "Offer and request interactions between users.",
                            "ownership_rationale": "Owns buyer-seller communication artifacts.",
                            "included_capabilities": [],
                            "excluded_capabilities": [],
                            "primary_entities": ["Offer"],
                            "allowed_dependencies": ["IdentityAndAccess", "ProductCatalog"],
                            "evidence_ids": ["E0004"],
                        },
                    ]
                },
                "analyze_context_IdentityAndAccess": {
                    "context": "IdentityAndAccess",
                    "description": "Registration and account access.",
                    "actors": [
                        {
                            "name": "NonRegisteredUsers",
                            "description": "Visitor who can register and browse.",
                            "evidence_ids": ["E0001"],
                        }
                    ],
                    "capabilities": [
                        {
                            "name": "RegisterAccount",
                            "description": "Register an account.",
                            "actor": "NonRegisteredUsers",
                            "evidence_ids": ["E0001"],
                        }
                    ],
                    "aggregate_roots": [
                        {
                            "name": "User",
                            "description": "Identity aggregate.",
                            "evidence_ids": ["E0001"],
                        }
                    ],
                    "entities": [
                        {
                            "name": "User",
                            "description": "Platform user.",
                            "evidence_ids": ["E0001"],
                        }
                    ],
                    "value_objects": [
                        {
                            "name": "ContactInformation",
                            "description": "Email and phone details.",
                            "attributes": ["email", "contactNumber"],
                            "evidence_ids": ["E0001"],
                        }
                    ],
                    "business_rules": [
                        {
                            "text": "Users must verify email before entering the platform.",
                            "category": "security",
                            "evidence_ids": ["E0001"],
                        }
                    ],
                    "domain_events": [
                        {
                            "name": "UserRegistered",
                            "description": "Registration completed.",
                            "evidence_ids": ["E0001"],
                        }
                    ],
                    "domain_services": [],
                    "allowed_dependencies": [],
                    "evidence_ids": ["E0001"],
                },
                "analyze_context_ProductCatalog": {
                    "context": "ProductCatalog",
                    "description": "Searchable product information.",
                    "actors": [
                        {
                            "name": "NonRegisteredUsers",
                            "description": "Visitor who can search products.",
                            "evidence_ids": ["E0002"],
                        }
                    ],
                    "capabilities": [
                        {
                            "name": "SearchProducts",
                            "description": "Search products by category or keywords.",
                            "actor": "NonRegisteredUsers",
                            "evidence_ids": ["E0002"],
                        },
                        {
                            "name": "ViewProductDetails",
                            "description": "View seller and product details.",
                            "actor": "NonRegisteredUsers",
                            "evidence_ids": ["E0002"],
                        }
                    ],
                    "aggregate_roots": [
                        {
                            "name": "Product",
                            "description": "Product information aggregate.",
                            "evidence_ids": ["E0002"],
                        }
                    ],
                    "entities": [
                        {
                            "name": "Product",
                            "description": "Sellable item.",
                            "evidence_ids": ["E0002"],
                        },
                        {
                            "name": "Category",
                            "description": "Product grouping.",
                            "evidence_ids": ["E0002"],
                        }
                    ],
                    "value_objects": [
                        {
                            "name": "Price",
                            "description": "Product price.",
                            "attributes": ["amount", "currency"],
                            "evidence_ids": ["E0002"],
                        }
                    ],
                    "business_rules": [],
                    "domain_events": [],
                    "domain_services": [],
                    "allowed_dependencies": [],
                    "evidence_ids": ["E0002"],
                },
                "analyze_context_SellerWorkspace": {
                    "context": "SellerWorkspace",
                    "description": "Seller-only workspace features.",
                    "actors": [
                        {
                            "name": "RegisteredUsers",
                            "description": "Authenticated seller.",
                            "evidence_ids": ["E0003"],
                        }
                    ],
                    "capabilities": [
                        {
                            "name": "AddProduct",
                            "description": "Add a new product listing.",
                            "actor": "RegisteredUsers",
                            "evidence_ids": ["E0003"],
                        },
                        {
                            "name": "ViewProductDashboard",
                            "description": "View seller dashboard.",
                            "actor": "RegisteredUsers",
                            "evidence_ids": ["E0003"],
                        }
                    ],
                    "aggregate_roots": [
                        {
                            "name": "Product",
                            "description": "Seller listing aggregate.",
                            "evidence_ids": ["E0003"],
                        }
                    ],
                    "entities": [
                        {
                            "name": "Product",
                            "description": "Sellable product listing.",
                            "evidence_ids": ["E0003"],
                        },
                        {
                            "name": "Dashboard",
                            "description": "Seller dashboard projection.",
                            "evidence_ids": ["E0003"],
                        }
                    ],
                    "value_objects": [
                        {
                            "name": "ProductDetails",
                            "description": "Condition and delivery details.",
                            "attributes": ["condition", "warranty", "homeDelivery"],
                            "evidence_ids": ["E0003"],
                        }
                    ],
                    "business_rules": [
                        {
                            "text": "Only the registered users can add products.",
                            "category": "security",
                            "evidence_ids": ["E0003"],
                        }
                    ],
                    "domain_events": [
                        {
                            "name": "ProductListed",
                            "description": "Product listing published.",
                            "evidence_ids": ["E0003"],
                        }
                    ],
                    "domain_services": [],
                    "allowed_dependencies": ["IdentityAndAccess", "ProductCatalog"],
                    "evidence_ids": ["E0003"],
                },
                "analyze_context_NegotiationAndCommunication": {
                    "context": "NegotiationAndCommunication",
                    "description": "Offer exchange between buyers and sellers.",
                    "actors": [
                        {
                            "name": "RegisteredUsers",
                            "description": "Authenticated user who can negotiate.",
                            "evidence_ids": ["E0004"],
                        }
                    ],
                    "capabilities": [
                        {
                            "name": "SendOffer",
                            "description": "Send offers to sellers.",
                            "actor": "RegisteredUsers",
                            "evidence_ids": ["E0004"],
                        }
                    ],
                    "aggregate_roots": [
                        {
                            "name": "Offer",
                            "description": "Offer exchange aggregate.",
                            "evidence_ids": ["E0004"],
                        }
                    ],
                    "entities": [
                        {
                            "name": "Offer",
                            "description": "Buyer and seller offer data.",
                            "evidence_ids": ["E0004"],
                        }
                    ],
                    "value_objects": [],
                    "business_rules": [],
                    "domain_events": [
                        {
                            "name": "OfferSubmitted",
                            "description": "An offer is submitted.",
                            "evidence_ids": ["E0004"],
                        }
                    ],
                    "domain_services": [],
                    "allowed_dependencies": ["IdentityAndAccess", "ProductCatalog"],
                    "evidence_ids": ["E0004"],
                },
                "Synthesizer": {
                    "schema_version": "2.0.0",
                    "project_name": "KinmailDomainModel",
                    "project_metadata": {
                        "version": "1.0.0",
                        "generated_at": "SET_BY_CODE",
                        "description": "Synthesized model."
                    },
                    "bounded_contexts": [
                        {
                            "context_name": "IdentityAndAccess",
                            "description": "Registration and access control.",
                            "allowed_dependencies": [],
                            "actors": [
                                {
                                    "name": "RegisteredUser",
                                    "description": "Authenticated user.",
                                    "confidence": 0.8,
                                    "evidence_ids": ["E0001"],
                                    "sources": []
                                }
                            ],
                            "capabilities": [
                                {
                                    "name": "RegisterAccount",
                                    "description": "Register an account.",
                                    "actor": "NonRegisteredUsers",
                                    "confidence": 0.8,
                                    "evidence_ids": ["E0001"],
                                    "sources": []
                                }
                            ],
                            "ubiquitous_language": {
                                "entities": [
                                    {
                                        "name": "User",
                                        "description": "Platform user.",
                                        "confidence": 0.8,
                                        "evidence_ids": ["E0001"],
                                        "sources": [],
                                        "synonyms_to_avoid": []
                                    }
                                ],
                                "value_objects": [
                                    {
                                        "name": "ContactInformation",
                                        "attributes": ["email", "contactNumber"],
                                        "description": "Email and phone details.",
                                        "confidence": 0.7,
                                        "evidence_ids": ["E0001"],
                                        "sources": []
                                    }
                                ],
                                "services": [],
                                "aggregates": [
                                    {
                                        "name": "User",
                                        "description": "Identity aggregate.",
                                        "confidence": 0.75,
                                        "evidence_ids": ["E0001"],
                                        "sources": []
                                    }
                                ],
                                "domain_events": [
                                    {
                                        "name": "UserRegistered",
                                        "description": "Registration completed.",
                                        "confidence": 0.7,
                                        "evidence_ids": ["E0001"],
                                        "sources": []
                                    }
                                ]
                            },
                            "business_rules": [
                                {
                                    "text": "Users must verify email before entering the platform.",
                                    "category": "security",
                                    "confidence": 0.8,
                                    "evidence_ids": ["E0001"],
                                    "sources": []
                                }
                            ],
                            "external_references": [],
                            "evidence_ids": ["E0001"],
                            "evidence": []
                        }
                    ],
                    "global_rules": {
                        "naming_convention": "PascalCase",
                        "banned_global_terms": ["Manager", "Util", "Helper", "Data", "Info"],
                        "cross_cutting_constraints": ["Only the registered users can add products."],
                        "assumptions": []
                    }
                },
            }
        )

        architect = DomainArchitect(provider=provider)
        architect.min_delay = 0
        analyses = architect.analyze_document(parsed.clean_text, parsed_documents=[parsed])
        model = architect.synthesize_final_model(analyses)

        assert model.schema_version == "2.0.0"
        assert len(model.bounded_contexts) >= 1
        first_context = model.bounded_contexts[0]
        assert len(first_context.actors) >= 1
        assert len(first_context.capabilities) >= 1
        assert len(first_context.business_rules) >= 1
        assert len(first_context.ubiquitous_language.aggregates) >= 1
        assert len(first_context.business_rules[0].sources) >= 1

        acceptance = {
            "actors": ["NonRegisteredUsers", "RegisteredUsers"],
            "entities": ["User", "Product", "Category", "Offer", "Dashboard"],
            "capabilities": [
                "RegisterAccount",
                "SearchProducts",
                "AddProduct",
                "ViewProductDashboard",
                "SendOffer",
            ],
            "constraints": [
                "Users must verify email before entering the platform.",
                "Only the registered users can add products.",
            ],
        }
        coverage = architect.evaluate_acceptance_coverage(model, acceptance)
        assert coverage["actors"]["coverage_percent"] == 100.0
        assert coverage["entities"]["coverage_percent"] == 100.0
        assert coverage["capabilities"]["coverage_percent"] == 100.0
        assert coverage["constraints"]["coverage_percent"] == 100.0

    def test_verifier_rejects_field_drop(self):
        from core.architect import DomainArchitect
        from core.schemas import (
            ContextAnalysis,
            ConstraintCandidate,
            DomainModel,
            EntityCandidate,
            ParsedSRSDocument,
            ParsedSection,
            ProjectMetadata,
            RequirementRecord,
            ScoutExtraction,
            RequirementSummary,
            EvidenceSpan,
        )

        architect = object.__new__(DomainArchitect)
        architect.stage_timings_ms = {}
        architect.last_parsed_documents = [
            ParsedSRSDocument(
                file_path="srs.docx",
                document_name="srs.docx",
                clean_text="Only the registered users can add products.",
                sections=[
                    ParsedSection(
                        section_id="SEC-001",
                        heading="Security Requirements",
                        category="security",
                        content="Only the registered users can add products.",
                        evidence_ids=["E0001"],
                    )
                ],
                requirements=[
                    RequirementRecord(
                        requirement_id="REQ-001",
                        category="security",
                        title="Only registered users can add products",
                        description="Only the registered users can add products.",
                        actor="RegisteredUsers",
                        section="Security Requirements",
                        evidence_ids=["E0001"],
                    )
                ],
                evidence_spans=[
                    EvidenceSpan(
                        evidence_id="E0001",
                        document="srs.docx",
                        section="Security Requirements",
                        excerpt="Only the registered users can add products.",
                        line=1,
                        requirement_id="REQ-001",
                    )
                ],
            )
        ]
        architect.last_scout_summary = ScoutExtraction(
            requirements=[
                RequirementSummary(
                    requirement_id="REQ-001",
                    title="Only registered users can add products",
                    category="security",
                    description="Only the registered users can add products.",
                    actor="RegisteredUsers",
                    evidence_ids=["E0001"],
                )
            ],
            actors=[],
            entities=[EntityCandidate(name="Product", description="Sellable product.", evidence_ids=["E0001"])],
            constraints=[ConstraintCandidate(text="Only the registered users can add products.", category="security", evidence_ids=["E0001"])],
            tables=[],
            capabilities=[],
            evidence_spans=[],
        )
        architect.last_context_analyses = [
            ContextAnalysis(
                context="SellerWorkspace",
                description="Seller features.",
                aggregate_roots=[EntityCandidate(name="Product", description="Aggregate.", evidence_ids=["E0001"])],
                entities=[EntityCandidate(name="Product", description="Sellable product.", evidence_ids=["E0001"])],
                business_rules=[ConstraintCandidate(text="Only the registered users can add products.", category="security", evidence_ids=["E0001"])],
                allowed_dependencies=[],
                evidence_ids=["E0001"],
            )
        ]

        model = DomainModel(
            project_name="TestDomainModel",
            project_metadata=ProjectMetadata(version="1.0.0", generated_at="2026-03-12", description="test"),
            bounded_contexts=[],
            global_rules=None,
        )

        report = architect.verify_model(model)
        assert report.passed is False
        assert any(item.startswith("context:SellerWorkspace") or item.startswith("aggregate:SellerWorkspace:Product") for item in report.missing_fields)


# =============================================================================
# CONFIG TESTS
# =============================================================================

class TestConfig:
    """Test configuration values."""
    
    def test_model_names(self):
        """Test model name configurations."""
        from config import AnalyzerConfig, ArchitectConfig
        
        # Validation uses flash-lite
        assert AnalyzerConfig.MODEL_NAME == "gemini-2.5-flash-lite"
        
        # Domain model generation uses flash
        assert ArchitectConfig.MODEL_NAME == "gemini-2.5-flash"
    
    def test_server_config(self):
        """Test server configuration."""
        from config import ServerConfig
        
        assert ServerConfig.HOST == "127.0.0.1"
        assert ServerConfig.PORT == 8000
    
    def test_rag_config(self):
        """Test RAG configuration."""
        from config import RAGConfig
        
        config = RAGConfig()
        
        assert config.CHUNK_SIZE > 0
        assert config.TOP_K > 0
        assert 0 <= config.MIN_RELEVANCE_SCORE <= 1


# =============================================================================
# EXPERIMENT SCORING TESTS
# =============================================================================

class TestExperimentScoring:
    def test_score_predictions_computes_metrics(self):
        from scoring import score_predictions

        summary = score_predictions(
            predictions=[
                {
                    "filename": "/tmp/sample.py",
                    "violations": [
                        {
                            "type": "SynonymViolation",
                            "message": "Class name 'ClientManager' contains a synonym 'client' for the term 'Customer'.",
                            "sources": [],
                        }
                    ],
                    "metrics": {
                        "validation_time_ms": 12.0,
                        "cost_usd": 0.001,
                        "parseable_outputs": 1,
                        "unparseable_outputs": 0,
                    },
                }
            ],
            ground_truth={
                "files": [
                    {
                        "path": "/tmp/sample.py",
                        "expected_violations": [
                            {"type": "SynonymViolation", "focus": "ClientManager"}
                        ],
                    }
                ]
            },
        )
        assert summary["micro"]["tp"] == 1
        assert summary["micro"]["precision"] == 1.0
        assert summary["parseable_output_rate_percent"] == 100.0


# =============================================================================
# FASTAPI APP TESTS (No server required)
# =============================================================================

class TestFastAPIRoutes:
    """Test FastAPI route definitions."""
    
    def test_app_has_health_endpoint(self):
        """Test /health endpoint exists."""
        from main import app
        routes = [route.path for route in app.routes]
        assert "/health" in routes
    
    def test_app_has_validate_endpoint(self):
        """Test /validate endpoint exists."""
        from main import app
        routes = [route.path for route in app.routes]
        assert "/validate" in routes
    
    def test_app_has_token_endpoints(self):
        """Test token tracking endpoints exist."""
        from main import app
        routes = [route.path for route in app.routes]
        assert "/tokens/stats" in routes
        assert "/tokens/summary" in routes
        assert "/tokens/reset" in routes
    
    def test_app_has_metrics_endpoints(self):
        """Test metrics endpoints exist."""
        from main import app
        routes = [route.path for route in app.routes]
        assert "/metrics/validation" in routes
        assert "/metrics/combined" in routes
        assert "/metrics/research" in routes
        assert "/metrics/export" in routes


# =============================================================================
# RUN CONFIGURATION
# =============================================================================

if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
