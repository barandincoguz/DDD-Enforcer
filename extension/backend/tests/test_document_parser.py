from pathlib import Path

import pytest
from docx import Document

from core.document_parser import SRSDocumentParser


def _write_simple_pdf(path: Path, pages: list[list[str]]) -> None:
    objects = []
    page_object_ids = []
    content_object_ids = []

    font_object_id = 3
    next_object_id = 4

    for _ in pages:
        page_object_ids.append(next_object_id)
        content_object_ids.append(next_object_id + 1)
        next_object_id += 2

    kids = " ".join(f"{page_id} 0 R" for page_id in page_object_ids)
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(
        f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode("latin-1")
    )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    for page_index, lines in enumerate(pages):
        page_id = page_object_ids[page_index]
        content_id = content_object_ids[page_index]

        stream_lines = ["BT", "/F1 12 Tf", "72 720 Td"]
        for line_index, line in enumerate(lines):
            escaped = (
                line.replace("\\", "\\\\")
                .replace("(", "\\(")
                .replace(")", "\\)")
            )
            if line_index > 0:
                stream_lines.append("0 -18 Td")
            stream_lines.append(f"({escaped}) Tj")
        stream_lines.append("ET")

        stream = "\n".join(stream_lines).encode("latin-1")
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {font_object_id} 0 R >> >> "
                f"/Contents {content_id} 0 R >>"
            ).encode("latin-1")
        )
        objects.append(
            f"<< /Length {len(stream)} >>\nstream\n".encode("latin-1")
            + stream
            + b"\nendstream"
        )

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]

    for object_id, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{object_id} 0 obj\n".encode("latin-1"))
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("latin-1"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))

    pdf.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\n".encode("latin-1")
    )
    pdf.extend(f"startxref\n{xref_offset}\n%%EOF".encode("latin-1"))
    path.write_bytes(bytes(pdf))


def test_parse_pdf_merges_wrapped_lines_and_stops_at_references(tmp_path):
    pdf_file = tmp_path / "requirements.pdf"
    _write_simple_pdf(
        pdf_file,
        [[
            "1. Requirements",
            "The system shall track",
            "orders for each customer.",
            "References",
            "This appendix should not be indexed.",
        ]],
    )

    content = SRSDocumentParser().parse_file(str(pdf_file))

    assert "The system shall track orders for each customer." in content
    assert "This appendix should not be indexed." not in content


def test_parse_docx_preserves_lists_and_tables(tmp_path):
    docx_file = tmp_path / "requirements.docx"
    document = Document()
    document.add_heading("Order Management", level=1)
    document.add_paragraph("The system shall store orders.")
    document.add_paragraph("Customer can cancel an order.", style="List Bullet")

    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Rule"
    table.rows[1].cells[0].text = "Status"
    table.rows[1].cells[1].text = "Required"

    document.save(docx_file)

    content = SRSDocumentParser().parse_file(str(docx_file))

    assert "Order Management" in content
    assert "- Customer can cancel an order." in content
    assert "Field | Rule" in content
    assert "Status | Required" in content


def test_parse_txt_supports_utf16_input(tmp_path):
    txt_file = tmp_path / "requirements.txt"
    txt_file.write_bytes("Odeme kaydi zorunludur.".encode("utf-16"))

    content = SRSDocumentParser().parse_file(str(txt_file))

    assert "Odeme kaydi zorunludur." in content


def test_parse_txt_does_not_truncate_regular_requirement_lines(tmp_path):
    txt_file = tmp_path / "references-prefix.txt"
    txt_file.write_text(
        "References to external systems shall be preserved.\n"
        "This rule remains important.",
        encoding="utf-8",
    )

    content = SRSDocumentParser().parse_file(str(txt_file))

    assert "References to external systems shall be preserved." in content
    assert "This rule remains important." in content


def test_parse_nonexistent_file_raises_file_not_found():
    parser = SRSDocumentParser()

    with pytest.raises(FileNotFoundError):
        parser.parse_file("/nonexistent/file.txt")


def test_parse_txt_truncates_at_kaynakca_heading(tmp_path):
    txt_file = tmp_path / "tr-references.txt"
    # 5 lines: bibliography at index 3 (>= midpoint 2). Should truncate.
    txt_file.write_text(
        "Müşteri siparişleri kayıt altına alınmalıdır.\n"
        "Sipariş iptalleri kayıt altına alınmalıdır.\n"
        "Sistem ödeme tahsil etmelidir.\n"
        "Kaynakça\n"
        "1. Smith, J. (2020). DDD in practice.\n",
        encoding="utf-8",
    )

    content = SRSDocumentParser().parse_file(str(txt_file))

    assert "Müşteri siparişleri" in content
    assert "Sistem ödeme tahsil etmelidir." in content
    assert "Smith, J." not in content


def test_parse_txt_truncates_at_kaynaklar_heading(tmp_path):
    txt_file = tmp_path / "tr-references-plural.txt"
    # 5 lines: bibliography at index 3 (>= midpoint 2). Should truncate.
    txt_file.write_text(
        "Sistem ödeme kaydını tutmalıdır.\n"
        "Müşteri profili oluşturulmalıdır.\n"
        "Sepet özeti hesaplanmalıdır.\n"
        "Kaynaklar\n"
        "1. Evans, E. (2003). Domain-Driven Design.\n",
        encoding="utf-8",
    )

    content = SRSDocumentParser().parse_file(str(txt_file))

    assert "Sistem ödeme" in content
    assert "Sepet özeti" in content
    assert "Evans, E." not in content


def test_parse_txt_truncates_at_numbered_top_level_references(tmp_path):
    txt_file = tmp_path / "numbered-references.txt"
    # 5 lines: bibliography at index 3 (>= midpoint 2). Should truncate.
    txt_file.write_text(
        "The system shall log all transactions.\n"
        "The system shall enforce idempotent submission.\n"
        "The system shall reject malformed payloads.\n"
        "5. References\n"
        "[1] IEEE 830-1998. Recommended Practice.\n",
        encoding="utf-8",
    )

    content = SRSDocumentParser().parse_file(str(txt_file))

    assert "log all transactions" in content
    assert "idempotent submission" in content
    assert "IEEE 830-1998" not in content


def test_parse_txt_does_not_truncate_at_nested_subsection_references_in_first_half(tmp_path):
    txt_file = tmp_path / "subsection-references.txt"
    # 8 lines: `3.4 References` at index 1 (< midpoint 4). Should NOT truncate.
    txt_file.write_text(
        "Section 3 — Functional Requirements.\n"
        "3.4 References\n"
        "The system shall accept ISO 8601 timestamps.\n"
        "The system shall reject unauthenticated requests.\n"
        "4. Non-functional Requirements\n"
        "The system shall respond within 200 ms.\n"
        "6. Glossary\n"
        "Term: aggregate root.\n",
        encoding="utf-8",
    )

    content = SRSDocumentParser().parse_file(str(txt_file))

    assert "3.4 References" in content
    assert "ISO 8601 timestamps" in content
    assert "respond within 200 ms" in content
    assert "Term: aggregate root." in content


def test_parse_txt_does_not_truncate_inline_kaynaklar_mention(tmp_path):
    txt_file = tmp_path / "inline-kaynaklar.txt"
    # 4 lines, inline use of "Kaynaklarımız" — never matches the regex.
    txt_file.write_text(
        "Sistem güvenliği önemlidir.\n"
        "Kaynaklarımız geniştir ve sistem buna göre tasarlanmalıdır.\n"
        "Veritabanı yedeklemeleri saatlik alınır.\n"
        "Sipariş yönetimi modülü güncellenmelidir.\n",
        encoding="utf-8",
    )

    content = SRSDocumentParser().parse_file(str(txt_file))

    assert "Kaynaklarımız geniştir" in content
    assert "Sipariş yönetimi" in content


def test_parse_txt_truncates_at_nested_5_1_references_at_end(tmp_path):
    txt_file = tmp_path / "appendix-references.txt"
    # 6 lines: `5.1 References` at index 4 (>= midpoint 3). Should truncate.
    # Verifies that the position guard does NOT block legitimate bibliography
    # subsections like `5.1 References` (appendix-style layout).
    txt_file.write_text(
        "The system shall persist orders.\n"
        "The system shall enforce price floors.\n"
        "5. Appendix\n"
        "Appendix A — error codes.\n"
        "5.1 References\n"
        "[1] Fowler, M. (2002). Enterprise Patterns.\n",
        encoding="utf-8",
    )

    content = SRSDocumentParser().parse_file(str(txt_file))

    assert "persist orders" in content
    assert "Appendix A" in content
    assert "Fowler, M." not in content


def test_parse_txt_truncates_at_references_with_trailing_colon(tmp_path):
    txt_file = tmp_path / "colon-references.txt"
    # 5 lines: `References:` at index 3 (>= midpoint 2). Should truncate.
    txt_file.write_text(
        "The system shall validate input.\n"
        "The system shall log audit trails.\n"
        "The system shall enforce role-based access.\n"
        "References:\n"
        "[1] OWASP ASVS 4.0.\n",
        encoding="utf-8",
    )

    content = SRSDocumentParser().parse_file(str(txt_file))

    assert "validate input" in content
    assert "role-based access" in content
    assert "OWASP" not in content


@pytest.mark.parametrize(
    "line,should_match",
    [
        # positive matches
        ("References", True),
        ("REFERENCES", True),
        ("Bibliography", True),
        ("BIBLIOGRAPHY", True),
        ("Kaynakça", True),
        ("KAYNAKÇA", True),
        ("Kaynaklar", True),
        ("KAYNAKLAR", True),
        ("# References", True),
        ("## Bibliography", True),
        ("5. References", True),
        ("5.1 References", True),
        ("5.1 Kaynakça", True),
        ("5.1.2 Kaynaklar", True),
        ("References:", True),
        ("Kaynaklar:", True),
        ("Kaynaklar：", True),  # U+FF1A fullwidth colon
        ("references", True),
        # negative — must NOT match
        ("References to external systems shall be preserved.", False),
        ("Kaynaklarımız geniştir.", False),
        ("Some References", False),
        ("3.4 References to Other Documents", False),
        ("3.4 Subsection References", False),
        ("Bibliography of authors:", False),
        ("", False),
        ("  ", False),
    ],
)
def test_reference_heading_pattern_direct_grammar(line, should_match):
    parser = SRSDocumentParser()
    match = parser.reference_heading_pattern.match(line.strip())
    assert bool(match) is should_match, f"Pattern mismatch for: {line!r}"


# ---------- WP-CORE-3 — Empty-input contract (F-3) ----------

import re

from core.document_parser import EmptySRSDocumentError


def test_parse_empty_txt_raises_empty_srs_document_error(tmp_path):
    empty = tmp_path / "empty.txt"
    empty.write_text("")
    parser = SRSDocumentParser()
    with pytest.raises(EmptySRSDocumentError):
        parser.parse_file(str(empty))


def test_parse_whitespace_only_txt_raises_empty_srs_document_error(tmp_path):
    ws = tmp_path / "ws.txt"
    ws.write_text("\n\n   \t\n")
    parser = SRSDocumentParser()
    with pytest.raises(EmptySRSDocumentError):
        parser.parse_file(str(ws))


def test_parse_empty_pdf_raises_empty_srs_document_error(tmp_path, monkeypatch):
    placeholder = tmp_path / "x.pdf"
    placeholder.write_bytes(b"%PDF-1.4\n%dummy\n")
    monkeypatch.setattr("core.document_parser.read_pdf", lambda p: "")
    parser = SRSDocumentParser()
    with pytest.raises(EmptySRSDocumentError):
        parser.parse_file(str(placeholder))


def test_parse_empty_docx_raises_empty_srs_document_error(tmp_path, monkeypatch):
    placeholder = tmp_path / "x.docx"
    placeholder.write_bytes(b"PK\x03\x04dummy")
    monkeypatch.setattr("core.document_parser.read_docx", lambda p: "")
    parser = SRSDocumentParser()
    with pytest.raises(EmptySRSDocumentError):
        parser.parse_file(str(placeholder))


def test_empty_srs_document_error_is_value_error(tmp_path):
    empty = tmp_path / "empty.txt"
    empty.write_text("")
    parser = SRSDocumentParser()
    with pytest.raises(ValueError):
        parser.parse_file(str(empty))


def test_parse_unsupported_extension_still_raises_plain_value_error(tmp_path):
    weird = tmp_path / "x.xyz"
    weird.write_text("anything")
    parser = SRSDocumentParser()
    with pytest.raises(ValueError, match="Unsupported file type") as exc_info:
        parser.parse_file(str(weird))
    assert not isinstance(exc_info.value, EmptySRSDocumentError)


def test_parse_non_empty_txt_returns_text_normally(tmp_path):
    good = tmp_path / "ok.txt"
    good.write_text("Hello world.\n")
    parser = SRSDocumentParser()
    assert parser.parse_file(str(good)) == "Hello world."


def test_parse_references_only_txt_raises_empty_srs_document_error(tmp_path):
    """Post-truncation-empty (WP-CORE-2 cross-integration) is caught."""
    only_refs = tmp_path / "only_refs.txt"
    only_refs.write_text("References\n")
    parser = SRSDocumentParser()
    with pytest.raises(EmptySRSDocumentError):
        parser.parse_file(str(only_refs))


def test_empty_srs_document_error_message_contains_file_path(tmp_path):
    empty = tmp_path / "empty.txt"
    empty.write_text("")
    parser = SRSDocumentParser()
    with pytest.raises(EmptySRSDocumentError, match=re.escape(str(empty))):
        parser.parse_file(str(empty))


def test_parse_nonexistent_file_still_raises_file_not_found_post_contract():
    parser = SRSDocumentParser()
    with pytest.raises(FileNotFoundError):
        parser.parse_file("/nonexistent/file.txt")


@pytest.mark.parametrize(
    "ext, patch_target",
    [
        (".pdf", "core.document_parser.read_pdf"),
        (".docx", "core.document_parser.read_docx"),
    ],
)
def test_each_binary_reader_empty_raises_empty_srs_document_error(
    tmp_path, monkeypatch, ext, patch_target
):
    placeholder = tmp_path / f"x{ext}"
    placeholder.write_bytes(b"placeholder bytes")
    monkeypatch.setattr(patch_target, lambda p: "")
    parser = SRSDocumentParser()
    with pytest.raises(EmptySRSDocumentError):
        parser.parse_file(str(placeholder))
